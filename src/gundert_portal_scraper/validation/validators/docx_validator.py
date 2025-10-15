"""DOCX content validator."""

import logging
from typing import Dict, List, Optional, Any
from pathlib import Path

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .. import BaseValidator, ValidationResult, ValidationIssue, ValidationSeverity

logger = logging.getLogger(__name__)


class DOCXValidator(BaseValidator):
    """Validator for DOCX (Microsoft Word) format."""
    
    def __init__(self):
        super().__init__()
        self.validator_name = "DOCXValidator"
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX validation will be limited.")
    
    def validate_content(self, content: str, format_type: str, options: Optional[Dict[str, Any]] = None) -> ValidationResult:
        """Validate DOCX content (from string is not supported for DOCX).
        
        Args:
            content: Not used for DOCX validation
            format_type: Should be 'docx'
            options: Optional validation options
            
        Returns:
            ValidationResult indicating string validation not supported
        """
        return ValidationResult(
            success=False,
            format_type=format_type,
            issues=[ValidationIssue(
                severity=ValidationSeverity.ERROR,
                code="DOCX_STRING_VALIDATION_UNSUPPORTED",
                message="DOCX validation from string content is not supported",
                suggestion="Use validate_file() method for DOCX files"
            )],
            metadata={'validator': self.validator_name, 'format_type': format_type}
        )
    
    def validate_file(self, file_path: Path, format_type: str, options: Optional[Dict[str, Any]] = None) -> ValidationResult:
        """Validate DOCX file.
        
        Args:
            file_path: Path to DOCX file to validate
            format_type: Should be 'docx'
            options: Optional validation options
            
        Returns:
            ValidationResult with validation issues
        """
        options = options or {}
        issues = []
        metadata = {
            'validator': self.validator_name,
            'format_type': format_type,
            'docx_available': DOCX_AVAILABLE,
            'file_path': str(file_path)
        }
        
        if not DOCX_AVAILABLE:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.CRITICAL,
                code="DOCX_LIBRARY_UNAVAILABLE",
                message="python-docx library not available for validation",
                suggestion="Install python-docx for DOCX validation"
            ))
            return ValidationResult(
                success=False,
                format_type=format_type,
                issues=issues,
                metadata=metadata
            )
        
        # Check file existence and extension
        if not file_path.exists():
            issues.append(ValidationIssue(
                severity=ValidationSeverity.CRITICAL,
                code="FILE_NOT_FOUND",
                message=f"DOCX file not found: {file_path}",
                location=str(file_path)
            ))
            return ValidationResult(
                success=False,
                format_type=format_type,
                issues=issues,
                metadata=metadata
            )
        
        if not file_path.suffix.lower() == '.docx':
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="UNEXPECTED_FILE_EXTENSION",
                message=f"File extension is '{file_path.suffix}', expected '.docx'",
                location=str(file_path)
            ))
        
        # Try to open and validate DOCX structure
        doc, doc_issues = self._load_document(file_path)
        issues.extend(doc_issues)
        
        if doc is not None:
            # Validate document structure
            issues.extend(self._validate_document_structure(doc))
            issues.extend(self._validate_content_quality(doc))
            issues.extend(self._check_malayalam_content(doc))
            issues.extend(self._validate_styles_and_formatting(doc))
            
            # Update metadata with document info
            metadata.update(self._get_document_metadata(doc, file_path))
        
        success = len([i for i in issues if i.severity in [ValidationSeverity.ERROR, ValidationSeverity.CRITICAL]]) == 0
        
        return ValidationResult(
            success=success,
            format_type=format_type,
            issues=issues,
            metadata=metadata
        )
    
    def _load_document(self, file_path: Path) -> tuple[Optional['Document'], List[ValidationIssue]]:
        """Load DOCX document and return any loading issues."""
        issues = []
        
        try:
            doc = Document(str(file_path))
            return doc, issues
            
        except PackageNotFoundError:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.CRITICAL,
                code="INVALID_DOCX_PACKAGE",
                message="File is not a valid DOCX package",
                location=str(file_path),
                suggestion="Ensure file is a valid Microsoft Word document"
            ))
            
        except Exception as e:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.CRITICAL,
                code="DOCX_LOAD_ERROR",
                message=f"Failed to load DOCX document: {e}",
                location=str(file_path),
                context={'exception': str(e)}
            ))
        
        return None, issues
    
    def _validate_document_structure(self, doc: 'Document') -> List[ValidationIssue]:
        """Validate basic document structure."""
        issues = []
        
        # Check if document has content
        if not doc.paragraphs:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="EMPTY_DOCUMENT",
                message="Document contains no paragraphs",
                suggestion="Add content to the document"
            ))
            return issues
        
        # Count content paragraphs (non-empty)
        content_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        if not content_paragraphs:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="NO_CONTENT_PARAGRAPHS",
                message="Document contains no paragraphs with text content",
                suggestion="Add text content to paragraphs"
            ))
        
        # Check for title/heading structure
        has_title = False
        for paragraph in doc.paragraphs[:5]:  # Check first 5 paragraphs
            if paragraph.style and 'title' in paragraph.style.name.lower():
                has_title = True
                break
            if paragraph.text.strip() and len(paragraph.text.strip()) > 10:
                # Assume first substantial paragraph is title-like
                has_title = True
                break
        
        if not has_title:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.INFO,
                code="NO_CLEAR_TITLE",
                message="Document doesn't appear to have a clear title",
                suggestion="Consider adding a title paragraph"
            ))
        
        return issues
    
    def _validate_content_quality(self, doc: 'Document') -> List[ValidationIssue]:
        """Validate content quality and structure."""
        issues = []
        
        total_text = '\n'.join(p.text for p in doc.paragraphs)
        
        # Check text length
        if len(total_text.strip()) < 100:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="MINIMAL_CONTENT",
                message=f"Document has minimal content ({len(total_text)} characters)",
                suggestion="Verify document contains expected content"
            ))
        
        # Check for very long paragraphs (might indicate formatting issues)
        for i, paragraph in enumerate(doc.paragraphs):
            if len(paragraph.text) > 2000:
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.INFO,
                    code="VERY_LONG_PARAGRAPH",
                    message=f"Paragraph {i+1} is very long ({len(paragraph.text)} characters)",
                    location=f"paragraph {i+1}",
                    suggestion="Consider breaking into smaller paragraphs"
                ))
        
        # Check for repeated content (might indicate processing error)
        paragraph_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if len(paragraph_texts) != len(set(paragraph_texts)):
            duplicate_count = len(paragraph_texts) - len(set(paragraph_texts))
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="DUPLICATE_PARAGRAPHS",
                message=f"Found {duplicate_count} duplicate paragraphs",
                suggestion="Review content for unintended duplication"
            ))
        
        return issues
    
    def _check_malayalam_content(self, doc: 'Document') -> List[ValidationIssue]:
        """Check for Malayalam content and encoding issues."""
        issues = []
        
        # Extract all text
        all_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                all_text.append(paragraph.text)
        
        combined_text = ' '.join(all_text)
        
        # Check for Malayalam characters
        malayalam_range = range(0x0D00, 0x0D80)
        has_malayalam = any(ord(char) in malayalam_range for char in combined_text)
        
        if not has_malayalam:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.WARNING,
                code="NO_MALAYALAM_CONTENT",
                message="No Malayalam characters detected in document",
                suggestion="Verify Malayalam text is properly encoded"
            ))
        else:
            # Check for potential encoding issues
            # Look for common encoding problem indicators
            problem_chars = ['�', '?']  # Replacement characters
            if any(char in combined_text for char in problem_chars):
                issues.append(ValidationIssue(
                    severity=ValidationSeverity.WARNING,
                    code="POTENTIAL_ENCODING_ISSUES",
                    message="Found replacement characters that may indicate encoding problems",
                    suggestion="Review document for character encoding issues"
                ))
            
            # Check Malayalam character distribution
            malayalam_chars = sum(1 for char in combined_text if ord(char) in malayalam_range)
            total_chars = len(combined_text.replace(' ', ''))
            
            if total_chars > 0:
                malayalam_percentage = (malayalam_chars / total_chars) * 100
                if malayalam_percentage < 10:
                    issues.append(ValidationIssue(
                        severity=ValidationSeverity.INFO,
                        code="LOW_MALAYALAM_CONTENT",
                        message=f"Malayalam characters make up only {malayalam_percentage:.1f}% of content",
                        context={'malayalam_percentage': malayalam_percentage}
                    ))
        
        return issues
    
    def _validate_styles_and_formatting(self, doc: 'Document') -> List[ValidationIssue]:
        """Validate document styles and formatting."""
        issues = []
        
        # Check style usage
        styles_used = set()
        for paragraph in doc.paragraphs:
            if paragraph.style:
                styles_used.add(paragraph.style.name)
        
        if len(styles_used) == 1 and 'Normal' in styles_used:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.INFO,
                code="MINIMAL_STYLE_USAGE",
                message="Document uses only Normal style",
                suggestion="Consider using heading and other styles for better structure"
            ))
        
        # Check for custom styles (indicator of specialized formatting)
        custom_styles = []
        for style_name in styles_used:
            if style_name not in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Title']:
                custom_styles.append(style_name)
        
        if custom_styles:
            issues.append(ValidationIssue(
                severity=ValidationSeverity.INFO,
                code="CUSTOM_STYLES_USED",
                message=f"Document uses custom styles: {', '.join(custom_styles[:5])}",
                context={'custom_styles': custom_styles}
            ))
        
        return issues
    
    def _get_document_metadata(self, doc: 'Document', file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata = {}
        
        # File metadata
        try:
            stat = file_path.stat()
            metadata['file_size_bytes'] = stat.st_size
            metadata['file_modification_time'] = stat.st_mtime
        except:
            pass
        
        # Document structure
        metadata['total_paragraphs'] = len(doc.paragraphs)
        metadata['content_paragraphs'] = len([p for p in doc.paragraphs if p.text.strip()])
        
        # Text statistics
        all_text = ' '.join(p.text for p in doc.paragraphs)
        metadata['total_characters'] = len(all_text)
        metadata['total_words'] = len(all_text.split())
        metadata['total_lines'] = len([p for p in doc.paragraphs if p.text])
        
        # Style information
        styles_used = set(p.style.name for p in doc.paragraphs if p.style)
        metadata['styles_used'] = list(styles_used)
        metadata['style_count'] = len(styles_used)
        
        # Document properties
        try:
            core_props = doc.core_properties
            metadata['document_title'] = core_props.title
            metadata['document_author'] = core_props.author
            metadata['document_subject'] = core_props.subject
            metadata['creation_date'] = core_props.created.isoformat() if core_props.created else None
            metadata['modification_date'] = core_props.modified.isoformat() if core_props.modified else None
        except:
            pass
        
        # Malayalam content statistics
        malayalam_range = range(0x0D00, 0x0D80)
        malayalam_chars = sum(1 for char in all_text if ord(char) in malayalam_range)
        metadata['malayalam_character_count'] = malayalam_chars
        
        if len(all_text) > 0:
            metadata['malayalam_percentage'] = (malayalam_chars / len(all_text)) * 100
        
        return metadata