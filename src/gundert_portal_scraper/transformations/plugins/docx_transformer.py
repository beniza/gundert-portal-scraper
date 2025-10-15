"""DOCX transformer for creating Word documents with preserved formatting."""

import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..framework import BaseTransformer, TransformationResult, LineMapping
from ...storage.schemas import BookStorage, PageContent
from ...core.exceptions import TransformationError

logger = logging.getLogger(__name__)


class DOCXTransformer(BaseTransformer):
    """Transforms content to DOCX format with preserved line mappings."""
    
    def __init__(self):
        super().__init__()
        self.output_format = "docx"
        self.supported_content_types = ["all"]
        self.version = "1.0"
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX transformation will not work.")
    
    def transform(self, book_storage: BookStorage, output_path: Optional[Path] = None, 
                 options: Dict[str, Any] = None) -> TransformationResult:
        """Transform book content to DOCX format.
        
        Args:
            book_storage: Source book storage
            output_path: Optional output file path
            options: Transformation options
            
        Returns:
            TransformationResult with DOCX document
        """
        if not DOCX_AVAILABLE:
            return TransformationResult(
                success=False,
                output_format=self.output_format,
                errors=["python-docx package not available. Install with: pip install python-docx"]
            )
        
        options = options or {}
        self.line_mappings = LineMapping()
        
        try:
            # Create DOCX document
            doc = self._create_document(book_storage, options)
            
            # Save to file if path provided
            file_path = None
            if output_path:
                output_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(output_path))
                file_path = str(output_path)
                logger.info(f"DOCX document saved to {output_path}")
            
            # Prepare metadata
            metadata = {
                'book_id': book_storage.book_metadata.book_id,
                'transformation_date': datetime.now().isoformat(),
                'source_pages': len(book_storage.pages),
                'total_paragraphs': len(doc.paragraphs),
                'options': options
            }
            
            return TransformationResult(
                success=True,
                output_format=self.output_format,
                content="[DOCX Document Created]",
                file_path=file_path,
                line_mappings=self.line_mappings,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"DOCX transformation failed: {e}")
            return TransformationResult(
                success=False,
                output_format=self.output_format,
                errors=[str(e)]
            )
    
    def _create_document(self, book_storage: BookStorage, options: Dict[str, Any]) -> 'Document':
        """Create DOCX document from book storage.
        
        Args:
            book_storage: Source book storage
            options: Transformation options
            
        Returns:
            python-docx Document object
        """
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = getattr(book_storage.book_metadata, 'title', 'Untitled')
        doc.core_properties.author = "Gundert Portal Scraper"
        doc.core_properties.subject = f"Book ID: {book_storage.book_metadata.book_id}"
        doc.core_properties.created = datetime.now()
        
        # Configure styles
        self._setup_styles(doc, options)
        
        # Add title
        self._add_title(doc, book_storage.book_metadata, options)
        
        # Add content from pages
        self._add_content(doc, book_storage, options)
        
        # Add appendix with metadata
        if options.get('include_metadata', True):
            self._add_metadata_appendix(doc, book_storage)
        
        return doc
    
    def _setup_styles(self, doc: 'Document', options: Dict[str, Any]):
        """Setup document styles.
        
        Args:
            doc: Document object
            options: Style options
        """
        styles = doc.styles
        
        # Configure Normal style for Malayalam text
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = options.get('font_name', 'Noto Sans Malayalam')
        normal_font.size = Pt(options.get('font_size', 12))
        
        # Create custom title style (avoid conflict with existing "Book Title" character style)
        try:
            title_style = styles.add_style('Document Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = options.get('title_font', 'Noto Sans Malayalam')
            title_font.size = Pt(options.get('title_size', 18))
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        except ValueError:
            # Style already exists
            pass
        
        # Create custom page header style (use existing "Header" style as base)
        try:
            page_style = styles.add_style('Page Section', WD_STYLE_TYPE.PARAGRAPH)
            page_font = page_style.font
            page_font.name = options.get('header_font', 'Arial')
            page_font.size = Pt(options.get('header_size', 10))
            page_font.bold = True
            page_style.paragraph_format.space_before = Pt(6)
            page_style.paragraph_format.space_after = Pt(3)
        except ValueError:
            # Style already exists
            pass
        
        # Create verse style for biblical content (use unique name)
        try:
            verse_style = styles.add_style('Bible Verse', WD_STYLE_TYPE.PARAGRAPH)
            verse_font = verse_style.font
            verse_font.name = options.get('font_name', 'Noto Sans Malayalam')
            verse_font.size = Pt(options.get('font_size', 12))
            verse_style.paragraph_format.left_indent = Inches(0.25)
            verse_style.paragraph_format.space_after = Pt(3)
        except ValueError:
            # Style already exists
            pass
    
    def _add_title(self, doc: 'Document', book_metadata, options: Dict[str, Any]):
        """Add document title.
        
        Args:
            doc: Document object
            book_metadata: Book metadata
            options: Title options
        """
        title = getattr(book_metadata, 'title', f"Book {book_metadata.book_id}")
        
        title_para = doc.add_paragraph()
        try:
            title_para.style = 'Document Title'
        except (ValueError, KeyError):
            # Fall back to Title style or Normal
            try:
                title_para.style = 'Title'
            except (ValueError, KeyError):
                title_para.style = 'Normal'
        title_para.add_run(title)
        
        # Add subtitle with book ID
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(f"Book ID: {book_metadata.book_id}")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(10)
        
        # Add source URL if available
        if hasattr(book_metadata, 'source_url') and book_metadata.source_url:
            source_para = doc.add_paragraph()
            source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source_run = source_para.add_run(f"Source: {book_metadata.source_url}")
            source_run.font.size = Pt(9)
            source_run.italic = True
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_content(self, doc: 'Document', book_storage: BookStorage, options: Dict[str, Any]):
        """Add content from book pages.
        
        Args:
            doc: Document object
            book_storage: Book storage
            options: Content options
        """
        include_page_headers = options.get('include_page_headers', True)
        preserve_line_breaks = options.get('preserve_line_breaks', True)
        detect_verses = options.get('detect_verses', True)
        
        # Verse detection patterns
        verse_patterns = [
            r'^\d+\.',  # "1.", "2.", etc.
            r'^\d+\s+',  # "1 ", "2 ", etc.
            r'വാക്യം\s*\d+',  # Verse in Malayalam
        ]
        
        for page_num, page in enumerate(book_storage.pages, 1):
            if not page.transcript_info.get('available'):
                continue
            
            # Add page header
            if include_page_headers:
                page_header = doc.add_paragraph()
                try:
                    page_header.style = 'Page Section'
                except (ValueError, KeyError):
                    # Fall back to Header style or Normal
                    try:
                        page_header.style = 'Header'
                    except (ValueError, KeyError):
                        page_header.style = 'Normal'
                page_header.add_run(f"Page {page_num}")
            
            # Get page content
            content_lines = self._get_page_lines(page)
            
            if not content_lines:
                # Empty page
                empty_para = doc.add_paragraph()
                empty_run = empty_para.add_run("[Empty page]")
                empty_run.italic = True
                empty_run.font.size = Pt(10)
                doc.add_paragraph()  # Add spacing
                continue
            
            # Process content lines
            for line_num, line in enumerate(content_lines, 1):
                line = line.strip()
                if not line:
                    if preserve_line_breaks:
                        doc.add_paragraph()
                    continue
                
                # Determine paragraph style
                is_verse = False
                if detect_verses:
                    for pattern in verse_patterns:
                        if re.match(pattern, line):
                            is_verse = True
                            break
                
                # Add paragraph
                para = doc.add_paragraph()
                if is_verse:
                    try:
                        para.style = 'Bible Verse'
                    except (ValueError, KeyError):
                        # Fall back to Body Text style or Normal
                        try:
                            para.style = 'Body Text'
                        except (ValueError, KeyError):
                            para.style = 'Normal'
                
                para.add_run(line)
                
                # Add line mapping
                para_index = len(doc.paragraphs) - 1
                self.line_mappings.add_mapping(
                    original_page=page_num,
                    original_line=line_num,
                    transformed_location=f"paragraph_{para_index}",
                    context={
                        'type': 'verse' if is_verse else 'content',
                        'page': page_num,
                        'paragraph_index': para_index
                    }
                )
            
            # Add page break (except for last page)
            if page_num < len(book_storage.pages) and options.get('page_breaks', False):
                doc.add_page_break()
            else:
                doc.add_paragraph()  # Add spacing between pages
    
    def _get_page_lines(self, page: PageContent) -> List[str]:
        """Extract text lines from page content.
        
        Args:
            page: Page content object
            
        Returns:
            List of text lines
        """
        lines = []
        
        # Get text content from transcript_info
        if page.transcript_info.get('available') and page.transcript_info.get('transcript_text'):
            lines = page.transcript_info['transcript_text'].split('\n')
        
        if not lines:
            return []
        
        # Clean and filter lines
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line and not self._is_metadata_line(line):
                cleaned_lines.append(line)
        
        return cleaned_lines
    
    def _is_metadata_line(self, line: str) -> bool:
        """Check if line contains metadata rather than content.
        
        Args:
            line: Line to check
            
        Returns:
            True if line appears to be metadata
        """
        metadata_indicators = [
            'page', 'പേജ്',  # Page indicators
            'source:', 'url:', 'http', 'www.',  # URL indicators
            '©', 'copyright', 'all rights reserved'  # Copyright indicators
        ]
        
        line_lower = line.lower()
        return any(indicator in line_lower for indicator in metadata_indicators)
    
    def _add_metadata_appendix(self, doc: 'Document', book_storage: BookStorage):
        """Add metadata appendix to document.
        
        Args:
            doc: Document object
            book_storage: Book storage
        """
        # Add page break
        doc.add_page_break()
        
        # Add appendix title
        appendix_title = doc.add_paragraph()
        try:
            appendix_title.style = 'Document Title'
        except (ValueError, KeyError):
            try:
                appendix_title.style = 'Heading 1'
            except (ValueError, KeyError):
                appendix_title.style = 'Normal'
        appendix_title.add_run("Document Metadata")
        
        # Add book metadata
        doc.add_paragraph().add_run("Book Information:").bold = True
        
        metadata = book_storage.book_metadata
        metadata_items = [
            ("Book ID", metadata.book_id),
            ("Title", getattr(metadata, 'title', 'N/A')),
            ("Language", getattr(metadata, 'language', 'Malayalam')),
            ("Content Type", getattr(metadata, 'content_type', 'N/A')),
            ("Source URL", getattr(metadata, 'source_url', 'N/A')),
        ]
        
        for label, value in metadata_items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))
        
        # Add page statistics
        doc.add_paragraph()
        doc.add_paragraph().add_run("Content Statistics:").bold = True
        
        total_pages = len(book_storage.pages)
        pages_with_content = sum(1 for p in book_storage.pages if p.transcript_info.get('available'))
        
        stats_items = [
            ("Total Pages", total_pages),
            ("Pages with Content", pages_with_content),
            ("Total Paragraphs", len(doc.paragraphs)),
            ("Line Mappings", len(self.line_mappings.mappings))
        ]
        
        for label, value in stats_items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))
        
        # Add transformation info
        doc.add_paragraph()
        doc.add_paragraph().add_run("Transformation Information:").bold = True
        
        transform_items = [
            ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Transformer", "DOCXTransformer v1.0"),
            ("Source", "Gundert Portal Scraper")
        ]
        
        for label, value in transform_items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))
    
    def validate_input(self, book_storage: BookStorage) -> List[str]:
        """Validate input for DOCX transformation.
        
        Args:
            book_storage: Book storage to validate
            
        Returns:
            List of validation errors
        """
        errors = super().validate_input(book_storage)
        
        if not DOCX_AVAILABLE:
            errors.append("python-docx package not available")
        
        return errors